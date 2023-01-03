from docx.enum.style import WD_STYLE_TYPE
from docx.table import Table

import docx
from docx.parts.document import Document
from docx.section import Section
from docx.text.paragraph import Paragraph
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.text.run import Run
from docx.enum.dml import MSO_COLOR_TYPE, MSO_THEME_COLOR_INDEX


class DocxStylesheet:
    def set_styles(self, doc: Document):
        styles = doc.styles
        heading_1 = styles['Heading 1']
        heading_1.font.color.rgb = None
        heading_1.font.all_caps = True
        heading_1.font.bold = True
        heading_1.font.size = Pt(14)
        heading_1.font.name = 'Times New Roman'
        heading_1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        heading_2 = styles['Heading 2']
        heading_2.font.color.rgb = None
        heading_2.font.bold = True
        heading_2.font.size = Pt(14)
        heading_2.paragraph_format.first_line_indent = Cm(1.25)
        heading_2.font.name = 'Times New Roman'
        heading_3 = styles['Heading 3']
        heading_3.font.color.rgb = None
        heading_3.font.bold = True
        heading_3.font.size = Pt(14)
        heading_3.paragraph_format.first_line_indent = Cm(1.25)
        heading_3.font.name = 'Times New Roman'

    def section_style(self, section: Section) -> Section:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        return section

    def paragraph_style(self, paragraph: Paragraph) -> Paragraph:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Cm(1.25)
        paragraph_format.line_spacing = 1.5
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        return paragraph

    def default_text_run(self, run: Run) -> Run:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.color.rgb = None
        return run

    def italic_run(self, run: Run) -> Run:
        run.font.italic = True
        return run

    def bold_run(self, run: Run) -> Run:
        run.font.bold = True
        return run

    def code_run(self, run: Run) -> Run:
        run.font.name = 'Courier New'
        run.font.size = Pt(12)
        return run

    def strike_run(self, run: Run) -> Run:
        run.font.strike = True
        return run

    def listing_style(self, paragraph: Paragraph) -> Paragraph:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format.line_spacing = 1.0
        [self.code_run(run) for run in paragraph.runs]
        return paragraph

    def all_caps_run(self, run: Run) -> Run:
        run.font.all_caps = True
        return run

    def listing_header_style(self, paragraph: Paragraph) -> Paragraph:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format.line_spacing = 1.0
        paragraph.style.font.name = 'Times New Roman'
        paragraph.style.font.size = Pt(14)
        paragraph.style.font.color.rgb = None
        return paragraph

    def table_header_style(self, paragraph: Paragraph) -> Paragraph:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph_format.line_spacing = 1.0
        paragraph.style.font.name = 'Times New Roman'
        paragraph.style.font.size = Pt(14)
        paragraph.style.font.color.rgb = None
        return paragraph


class DocxTreeVisitor:

    def __init__(self):
        self.doc_name = 'docx/article.docx'
        self.doc: Document = docx.Document()
        self.section: Section = self.doc.sections[-1]
        self.paragraph: Paragraph = None
        self.table: Table = None
        self.table_coords = [0, 0]
        self.header_numbers = {2: 0, 3: 0}
        self.listing_counter = 0
        self.image_counter = 0
        self.table_counter = 0

        self.stylesheet = DocxStylesheet()
        self.stylesheet.set_styles(self.doc)
        self.stylesheet.section_style(self.section)

    def visit(self, node):
        if type(node) == list:
            return sum([self.visit(node_item) for node_item in node], start=[])
        if type(node) != dict:
            return
        if "type" not in node:
            return
        node_type = node["type"]
        visit_method_name = f"_visit_{node_type}"
        if hasattr(self, visit_method_name) and callable(
                getattr(self, visit_method_name)):
            res = getattr(self, visit_method_name)(node)
            return res
        else:
            print(f'no such visit method {visit_method_name}')

    def _visit_heading(self, node):
        level = node["attrs"]["level"]
        header_number = ''
        if level == 2:
            self.header_numbers[2] += 1
            self.header_numbers[3] = 0
            header_number = f'{self.header_numbers[2]} '
        elif level == 3:
            self.header_numbers[3] += 1
            header_number = f'{self.header_numbers[2]}.{self.header_numbers[3]} '
        self.paragraph = self.doc.add_heading('', level=level)
        run = self.paragraph.add_run(header_number)
        self.stylesheet.default_text_run(run)
        self.paragraph.style = self.doc.styles[f'Heading {level}']
        self.visit(node['children'])
        return [self.paragraph]

    def _visit_text(self, node):
        if not node['raw'].endswith(' '):
            node['raw'] = node['raw'] + ' '
        run = self.paragraph.add_run(node['raw'])
        self.stylesheet.default_text_run(run)
        return [run]

    def _visit_blank_line(self, node):
        run = self.paragraph.add_run(' ')
        return [run]

    def _visit_paragraph(self, node):
        self.paragraph = self.doc.add_paragraph()
        self.stylesheet.paragraph_style(self.paragraph)
        self.visit(node['children'])
        return [self.paragraph]

    def _visit_block_code(self, node):
        self.listing_counter += 1
        listing_header_text = f'Листинг {self.listing_counter} – Название листинга'
        self.paragraph = self.doc.add_paragraph(listing_header_text)
        self.stylesheet.listing_header_style(self.paragraph)
        self.paragraph = self.doc.add_paragraph(node['raw'])
        self.stylesheet.listing_style(self.paragraph)
        return [self.paragraph]

    def _visit_softbreak(self, node):
        return []

    def _visit_link(self, node):
        return []

    def _visit_codespan(self, node):
        run = self.paragraph.add_run(node['raw'])
        self.stylesheet.code_run(run)
        return [run]

    def _visit_strikethrough(self, node):
        runs = self.visit(node['children'])
        [self.stylesheet.strike_run(run) for run in runs]
        return runs

    def _visit_table(self, node):
        self.table_counter += 1
        table_name = f'Таблица {self.table_counter} – Название таблицы'
        self.paragraph = self.doc.add_paragraph(table_name)
        self.stylesheet.table_header_style(self.paragraph)
        table_width = len(node['children'][0]['children'])
        table_height = len(node['children'][1]['children']) + 1
        self.table = self.doc.add_table(rows=table_height, cols=table_width)
        self.table.style = self.doc.styles['Table Grid']
        self.table_coords = [0, 0]
        [self.visit(child) for child in node['children']]
        return [self.table]

    def _visit_table_head(self, node):
        res = sum([self.visit(child) for child in node['children']], start=[])
        self.table_coords[0] += 1
        return res

    def _visit_table_body(self, node):
        res = sum([self.visit(child) for child in node['children']], start=[])
        return [res]

    def _visit_table_row(self, node):
        self.table_coords[1] = 0
        [self.visit(child) for child in node['children']]
        self.table_coords[0] += 1
        return []

    def _visit_table_cell(self, node):
        cell = self.table.cell(*self.table_coords)
        self.paragraph = cell.paragraphs[-1]
        [self.visit(child) for child in node['children']]
        self.table_coords[1] += 1
        return []

    def _visit_list(self, node):
        return []

    def _visit_list_item(self, node):
        return []

    def _visit_image(self, node):
        url = node['attrs']['url']
        self.paragraph = self.doc.add_paragraph()
        run = self.paragraph.add_run()
        picture = run.add_picture(url, width=Cm(15))
        self.paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.paragraph.paragraph_format.line_spacing = 1.0
        self.image_counter += 1
        self.paragraph = self.doc.add_paragraph(f'Рисунок {self.image_counter} – ')
        self.paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.paragraph.paragraph_format.line_spacing = 1.0
        children_res = sum([self.visit(child) for child in node['children']], start=[])
        return [self.paragraph]

    def _visit_block_text(self, node):
        return []

    def _visit_emphasis(self, node):
        children_res = self.visit(node['children'])
        [self.stylesheet.italic_run(run) for run in children_res]
        return children_res

    def _visit_strong(self, node):
        children_res = self.visit(node['children'])
        [self.stylesheet.bold_run(run) for run in children_res]
        return children_res

    def _visit_inline_math(self, node):
        self.paragraph = self.doc.add_paragraph(f"$${node['raw']}$$")
        # self.paragraph.style.font.math = True
        return [self.paragraph]
