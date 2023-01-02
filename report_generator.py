from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt
import json
import os
from pathlib import Path
from balaboba_request import generateDescription 

# Используемые файлы
TITUL_PAGE = "Titul_page.docx"
OUTPUT_PAGE = "output.docx"
CONFIG_JSON = "config.json"

# Используемые стили
HEADER_STYLE = "BoldHeader"
HEADER_LINK_STYLE = "BoldHeaderHyperlink"
CONTENT_STYLE = "Content"
CODE_STYLE = "Code"

def addTextParagraphToDocumentInStyle(text, document, style):
    p = document.add_paragraph(text)
    p.style = document.styles[style]

headers = [
    "Теоретическое введение",
    "Постановка задачи",
    "Программный код",
    "Вывод"
]

def contentStyle(styles, styleName):
    style = styles.add_style(styleName, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    style.paragraph_format.first_line_indent = Cm(1.25)
    return style

def boldHeaderStyle(styles, styleName):
    style = styles.add_style(styleName, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    style.font.size = Pt(16)
    style.font.bold = True
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return style

def codeStyle(styles, styleName):
    style = styles.add_style(styleName, WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Courier New'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1  #1.5
    style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return style


userStyles = {
    CONTENT_STYLE : contentStyle,
    HEADER_STYLE: boldHeaderStyle,
    HEADER_LINK_STYLE: boldHeaderStyle,
    CODE_STYLE : codeStyle
}


def defineAllStyles(styles):
    for styleName in userStyles:
        print(styleName)
        style = userStyles[styleName](styles, styleName)

def getConfigFile(config):
    with open(config, encoding='utf-8') as json_file:
            # encoding='utf-8' - для корректного отображения русских символов
        return json.load(json_file)

def writeCodeToDocument(task, document):
    path = Path(os.getcwd()).joinpath(task["folder"])
    tree = os.walk(path)

    for root, subdirs, files in tree:
        if subdirs == []:
            folder = root.split("\\")[-1]
            print(f'--\nfolder = {folder}')
            
            for filename in files:
                if (filename.split(".")[0] in task["include"] and filename.endswith(".java")):
                    filepath = os.path.join(root, filename)
                    print(f'\t- file {filename}')
                    print(filepath)
                    with open(filepath, encoding="utf-8") as file:
                        # encoding='utf-8' - для корректного отображения русских символов
                        code = file.read()
                        addTextParagraphToDocumentInStyle(code, document, CODE_STYLE)




def main():
    document = Document(TITUL_PAGE)
    
    # Определяем стили
    styles = document.styles
    defineAllStyles(styles)

    # Получаем конфигурационный файл
    config = getConfigFile(CONFIG_JSON)
    tasks = config["tasks"]

    # Заполняем документ
    i = 1
    for task in tasks:
        # header = f"Практическая работа №{i} «{task['topic']}»"
        header = task['topic']
        addTextParagraphToDocumentInStyle(header, document, HEADER_LINK_STYLE)

        # Теоретическое введение
        # addTextParagraphToDocumentInStyle(headers[0], document, HEADER_STYLE)
        # addTextParagraphToDocumentInStyle(generateDescription(task["topic"], "Введение"), document, CONTENT_STYLE)
        
        # Постановка задачи
        addTextParagraphToDocumentInStyle(headers[1], document, HEADER_STYLE) 
        addTextParagraphToDocumentInStyle(task["problem_statement"], document, CONTENT_STYLE)
        # Программный код
        addTextParagraphToDocumentInStyle(headers[2], document, HEADER_STYLE) 
        writeCodeToDocument(task, document)

        # Вывод
        # addTextParagraphToDocumentInStyle(headers[3], document, HEADER_STYLE)
        # addTextParagraphToDocumentInStyle(generateDescription(task["topic"], "Вывод"), document, CONTENT_STYLE)
        i+=1
    document.save(OUTPUT_PAGE)


main()


